from django.http import FileResponse 
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from docx import Document
from docx.shared import Inches
from mailmerge import MailMerge
from .models import PlaceHolder, UserProfile, PHChaptor0, PHChaptor1, PHChaptor2
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.shortcuts import render
from .forms import PlaceHolderForm
from django.contrib.auth.models import User
from .forms import *


# docx目录
class Index(ListView):
    model = PlaceHolder

    def get_queryset(self):
        return PlaceHolder.objects.filter(user=self.request.user)

class PlaceHolderDetail(DetailView):
    model = PlaceHolder

# 新建docx
class PlaceHolderCreate(CreateView):
    model = PlaceHolder
    form_class = PlaceHolderForm

    def form_valid(self, form):
        form.instance.user = self.request.user
        return super().form_valid(form)


def PlaceHolderCreateNew(request):
    if request.method == 'GET':
        ph = PlaceHolderForm()
        ph0 = PH0Form()
        ph1 = PH1Form()
        ph2 = PH2Form()
        return render(request, 'main/placeholder_form.html', {
            'ph': ph,
            'ph0': ph0,
            'ph1': ph1,
            'ph2': ph2
        })
    else:
        ph = PlaceHolderForm(request.POST, request.FILES)
        ph0 = PH0Form(request.POST, request.FILES)
        ph1 = PH1Form(request.POST, request.FILES)
        ph2 = PH2Form(request.POST, request.FILES)
        ph0m = ph0.save()
        ph1m = ph1.save()
        ph2m = ph2.save()
        phm = ph.save(commit=False)
        phm.chaptor0 = ph0m
        phm.chaptor1 = ph1m
        phm.chaptor2 = ph2m
        phm.user = request.user
        phm.save()
        return HttpResponseRedirect(reverse('list'))


def PlaceHolderUpdateNew(request, pk):
    phm = PlaceHolder.objects.get(pk=pk)
    if request.method == 'GET':
        ph = PlaceHolderForm(instance=phm)
        ph0 = PH0Form(instance=phm.chaptor0)
        ph1 = PH1Form(instance=phm.chaptor1)
        ph2 = PH2Form(instance=phm.chaptor2)
        return render(request, 'main/placeholder_form.html', {
            'ph': ph,
            'ph0': ph0,
            'ph1': ph1,
            'ph2': ph2
        })
    else:
        ph = PlaceHolderForm(request.POST, request.FILES, instance=phm)
        ph0 = PH0Form(request.POST, request.FILES, instance=phm.chaptor0)
        ph1 = PH1Form(request.POST, request.FILES, instance=phm.chaptor1)
        ph2 = PH2Form(request.POST, request.FILES, instance=phm.chaptor2)
        ph0m = ph0.save()
        ph1m = ph1.save()
        ph2m = ph2.save()
        phm.chaptor0 = ph0m
        phm.chaptor1 = ph1m
        phm.cpaptor2 = ph2m
        phm.save()
        return HttpResponseRedirect(reverse('list'))

# 修改docx
class PlaceHolderUpdate(UpdateView):
    model = PlaceHolder
    form_class = PlaceHolderForm

# 删除docx
class PlaceHolderDelete(DeleteView):
    model = PlaceHolder
    success_url = reverse_lazy('list')

# pk:primary key 创建docx
def makeDocx(pk):
    ph = PlaceHolder.objects.get(pk=pk)
    ch0 = PHChaptor0.objects.filter(pk=ph.chaptor0.pk).values()[0]
    ch1 = PHChaptor1.objects.filter(pk=ph.chaptor1.pk).values()[0]
    ch2 = PHChaptor2.objects.filter(pk=ph.chaptor2.pk).values()[0]
    obj = {**ch0, **ch1, **ch2}
    obj['h0'] = obj['d0']
    obj['h1'] = obj['d1']
    obj['h2'] = obj['d1_m']

    doc = MailMerge('temp')
    doc.merge(**obj)
    doc.write('result.docx')

    doc = Document('result.docx')
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == '@@@PP:i155@@@' or txt == '@@@PP:i156@@@':
            path = ''
            if txt == '@@@PP:i155@@@':
                path = ph.chaptor1.i155.path if ph.chaptor1.i155 else ''
            else:
                path = ph.chaptor1.i156.path if ph.chaptor1.i156 else ''
            p.clear()
            r = p.add_run()
            if path:
                #print(path)
                r.add_picture(path, width=Inches(5))

    doc.save(ph.filename + '.docx')
    return ph.filename + '.docx'


# 创建多个docx, request.GET['ids']为docx的primary key数组, json格式
def GetDocxList(request):
    import json
    import zipfile
    lst = json.loads(request.GET['ids'])
    with zipfile.ZipFile('result.zip', 'w') as f:
        for id in lst:
            fname = makeDocx(id)
            f.write(fname)
    return FileResponse(open('result.zip', 'rb'), as_attachment=False, filename='result.zip')

# 获取一个docx
def GetDocx(request, pk):
    fname = makeDocx(pk)
    return FileResponse(open(fname, 'rb'), as_attachment=True, filename=fname)


def UserLogin(request):
    if request.method == 'GET':
        return render(request, 'main/login.html', {})
    else:
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return HttpResponseRedirect(reverse('list'))
        else:
            context = {
                'username': username,
                'password': password,
            }
            return render(request, 'main/login.html', context)


def UserLogout(request):
    logout(request)
    return HttpResponseRedirect(reverse('login'))


def UserRegister(request):
    if request.method == 'GET':
        return render(request, 'main/register.html', {})
    else:
        username = request.POST['username']
        password = request.POST['password']
        sex = int(request.POST['sex']) == 0
        org = request.POST['org']
        try:
            user = User.objects.create_user(username=username, password=password)
            user.save()
            profile = UserProfile(user=user, org=org, sex=sex)
            profile.save()
        except:
            return render(request, 'main/register.html', {
                'username': username
            })
        return HttpResponseRedirect(reverse('list'))


def ChangePass(request):
    if request.method == 'GET':
        return render(request, 'main/changepass.html', {})
    else:
        password = request.POST['password']
        newpassword = request.POST['newpassword']
        if authenticate(username=request.user.username, password=password):
            request.user.set_password(newpassword)
            request.user.save()
            return HttpResponseRedirect(reverse('list'))
        return render(request, 'main/changepass.html', {
            'status': 'error'
        })
